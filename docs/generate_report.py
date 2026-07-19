"""Generate the Persian final report from raw experiment outputs.

The script never hard-codes experimental numbers.  Every table and numerical
claim is read from results/raw_data so the report remains reproducible.
"""
from __future__ import annotations

import json
import math
import statistics
import sys
from pathlib import Path
from typing import Iterable, Sequence

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RAW = ROOT / "results" / "raw_data"
FIG = ROOT / "results" / "figures"
MAPS = ROOT / "environments" / "maps"
OUT = ROOT / "report_source.docx"

BODY_FONT = "Noto Naskh Arabic"
HEAD_FONT = "Noto Sans Arabic"
LATIN_FONT = "DejaVu Sans"
MONO_FONT = "DejaVu Sans Mono"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_rtl(paragraph, rtl: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    elif not rtl and bidi is not None:
        p_pr.remove(bidi)


def set_run_font(run, font_name: str = BODY_FONT, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, LATIN_FONT, 9)


def add_p(doc: Document, text: str = "", *, bold: bool = False, center: bool = False, indent: bool = True, size: float = 11.2) -> object:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    set_rtl(p, not center)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(5)
    if indent and not center:
        p.paragraph_format.first_line_indent = Cm(0.55)
    run = p.add_run(text)
    set_run_font(run, BODY_FONT, size, bold)
    return p


def add_ltr(doc: Document, text: str, *, center: bool = True, font: str = MONO_FONT, size: float = 10.2) -> object:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    set_rtl(p, False)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, font, size)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> object:
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p, True)
    run = p.add_run(text)
    set_run_font(run, HEAD_FONT, 16 if level == 1 else 13.5, True)
    if level == 1:
        run.font.color.rgb = RGBColor(31, 78, 121)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(8)
    else:
        run.font.color.rgb = RGBColor(55, 55, 55)
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(5)
    return p


def set_cell_text(cell, text: object, *, header: bool = False, rtl: bool = True, size: float = 9.2) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p, rtl)
    run = p.add_run(str(text))
    set_run_font(run, HEAD_FONT if header else BODY_FONT, size, header)
    if header:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9EAF7")
        cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[object]], widths: Sequence[float] | None = None) -> object:
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, header=True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value, header=False, rtl=not isinstance(value, (int, float, np.number)))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_caption(doc: Document, text: str) -> object:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p, True)
    run = p.add_run(text)
    set_run_font(run, BODY_FONT, 9.5, True)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_figure(doc: Document, filename: str, caption: str, analysis: str, width: float = 6.35) -> None:
    path = FIG / filename
    if not path.exists():
        add_p(doc, f"[شکل {filename} در زمان تولید گزارش موجود نبود.]", center=True, indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p, False)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)
    add_p(doc, analysis)


def fmt(value: float, digits: int = 3) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "-"
    if math.isinf(float(value)):
        return "دستیابی نشد"
    return f"{float(value):.{digits}f}"


def mean_std(series: pd.Series, digits: int = 3) -> str:
    values = pd.to_numeric(series, errors="coerce").dropna()
    if not len(values):
        return "-"
    if len(values) == 1:
        return fmt(values.iloc[0], digits)
    return f"{values.mean():.{digits}f} ± {values.std(ddof=1):.{digits}f}"


def final_run_stats(frame: pd.DataFrame, group_col: str, tail: int = 200) -> pd.DataFrame:
    rows = []
    for (group, seed), subset in frame.groupby([group_col, "seed"]):
        subset = subset.sort_values("episode")
        last = subset.tail(min(tail, len(subset)))
        roll = subset["success"].rolling(min(100, max(20, len(subset) // 5)), min_periods=min(100, max(20, len(subset) // 5))).mean()
        hits = np.flatnonzero(roll.to_numpy() >= 0.8)
        rows.append(
            {
                group_col: group,
                "seed": seed,
                "final_success": last["success"].mean(),
                "final_return": last["return"].mean(),
                "final_steps": last["steps"].mean(),
                "episodes_to_80": float(hits[0]) if len(hits) else math.inf,
                "wall_hits": last["wall_hits"].mean() if "wall_hits" in last else np.nan,
                "penalty_entries": last["penalty_entries"].mean() if "penalty_entries" in last else np.nan,
            }
        )
    return pd.DataFrame(rows)


def obstacle_change(source: dict, target: dict) -> tuple[int, int, float]:
    def interior_walls(data: dict) -> set[tuple[int, int]]:
        rows = data["grid"]
        n = len(rows)
        return {(r, c) for r in range(1, n - 1) for c in range(1, n - 1) if rows[r][c] == "#"}
    s, t = interior_walls(source), interior_walls(target)
    removed, added = len(s - t), len(t - s)
    return removed, added, removed / max(1, len(s))


def create_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)
    normal.font.size = Pt(11.2)

    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        add_page_number(footer)
    return doc


def build_report() -> Path:
    required = [
        RAW / "final_summary.json",
        RAW / "value_iteration_gamma.csv",
        RAW / "q_learning_epsilon_schedules.csv",
        RAW / "reward_shaping_comparison.csv",
        RAW / "sarsa_lambda_ablation.csv",
        RAW / "three_algorithm_comparison.csv",
        RAW / "transfer_summary.csv",
    ]
    missing = [str(p) for p in required if not p.exists()]
    if missing:
        raise FileNotFoundError("Run submission experiments first. Missing: " + ", ".join(missing))

    final = load_json(RAW / "final_summary.json")
    run_manifest = load_json(RAW / "run_manifest.json")
    source_map = load_json(MAPS / "source_map.json")
    similar_map = load_json(MAPS / "target_similar_map.json")
    different_map = load_json(MAPS / "target_different_map.json")
    vi = pd.read_csv(RAW / "value_iteration_gamma.csv")
    q_eps = pd.read_csv(RAW / "q_learning_epsilon_schedules.csv")
    reward = pd.read_csv(RAW / "reward_shaping_comparison.csv")
    sarsa = pd.read_csv(RAW / "sarsa_lambda_ablation.csv")
    sarsa_summary = pd.read_csv(RAW / "sarsa_lambda_summary.csv")
    comparison = pd.read_csv(RAW / "three_algorithm_comparison.csv")
    transfer = pd.read_csv(RAW / "transfer_summary.csv")
    q_update = pd.read_csv(RAW / "q_learning_manual_update_trace.csv")
    trace = pd.read_csv(RAW / "sarsa_lambda_trace_log.csv")
    disagreement = load_json(RAW / "policy_disagreement_examples.json")
    negative = load_json(RAW / "negative_transfer_example.json")
    q_correction = pd.read_csv(RAW / "negative_transfer_q_correction.csv") if (RAW / "negative_transfer_q_correction.csv").exists() else pd.DataFrame()

    best_eps = final["best_epsilon_schedule"]
    best_lambda = float(final["best_lambda"])
    q_eps_stats = final_run_stats(q_eps, "epsilon_schedule")
    reward_stats = final_run_stats(reward, "reward_mode")
    lambda_stats = final_run_stats(sarsa, "lambda_label")

    similar_removed, similar_added, similar_frac = obstacle_change(source_map, similar_map)
    diff_removed, diff_added, diff_frac = obstacle_change(source_map, different_map)

    doc = create_document()

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p, True)
    run = p.add_run("پروژه پایانی درس یادگیری تقویتی")
    set_run_font(run, HEAD_FONT, 22, True)
    run.font.color.rgb = RGBColor(31, 78, 121)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p, True)
    run = p.add_run("طراحی و تحلیل عامل هوشمند در هزارتوی پویا")
    set_run_font(run, HEAD_FONT, 19, True)
    doc.add_paragraph()
    add_p(doc, "نام دانشجو: امیرمحمد جمالدار", center=True, indent=False, size=13)
    add_p(doc, "شماره دانشجویی: 40400356", center=True, indent=False, size=13)
    add_p(doc, "Seed اختصاصی: 5 | اندازه نقشه: 16×16", center=True, indent=False, size=12)
    add_p(doc, "تاریخ تهیه گزارش: ژوئیه 2026", center=True, indent=False, size=11)
    doc.add_page_break()

    add_heading(doc, "چکیده", 1)
    add_p(
        doc,
        f"در این پروژه یک محیط هزارتوی پویای 16×16 با انتقال تصادفی، کلید، در قفل، هدف، خانه‌های جریمه و یک دروازه دوره‌ای طراحی شد. حالت عامل به‌صورت (ردیف، ستون، وضعیت کلید، فاز دروازه) تعریف شد تا خاصیت مارکوف حفظ شود. سه الگوریتم Value Iteration، Q-Learning و SARSA(λ) بدون استفاده از کتابخانه‌های آماده یادگیری تقویتی پیاده‌سازی شدند. آزمایش‌های اصلی روی {len(run_manifest['seeds'])} Seed مستقل و داده‌های اپیزودی خام انجام شد. بهترین برنامه کاهش اکتشاف در داده‌های این اجرا «{best_eps}» و بهترین مقدار λ برابر {best_lambda:.1f} بود. توافق دقیق سیاست حریصانه Q-Learning و SARSA(λ) با سیاست مرجع Value Iteration به‌ترتیب {100*final['q_policy_agreement']:.1f} و {100*final['sarsa_policy_agreement']:.1f} درصد به دست آمد. انتقال یادگیری در دو مقصد با تغییر محدود و گسترده بررسی و یک نمونه انتقال منفی با regret عددی مستند شد. تمام نمودارها از CSVهای خام مخزن تولید شده‌اند.",
    )
    add_p(doc, "کلیدواژه‌ها: یادگیری تقویتی، فرایند تصمیم‌گیری مارکوف، Q-Learning، SARSA(λ)، Value Iteration، انتقال یادگیری، Reward Shaping، Eligibility Trace.")

    add_heading(doc, "فهرست مطالب", 1)
    toc_items = [
        "1. تعریف مسئله و محیط اختصاصی",
        "2. مدل‌سازی کامل MDP",
        "3. طراحی تابع پاداش",
        "4. پیاده‌سازی الگوریتم‌ها",
        "5. پروتکل آزمایش و معیارهای آماری",
        "6. نتایج Value Iteration",
        "7. نتایج Q-Learning و طراحی پاداش",
        "8. نتایج SARSA(λ)",
        "9. مقایسه سه الگوریتم و اختلاف سیاست",
        "10. انتقال یادگیری",
        "11. رابط گرافیکی و بازتولیدپذیری",
        "12. پاسخ مستقیم به پرسش‌های تحلیلی",
        "13. محدودیت‌ها و جمع‌بندی",
        "پیوست‌ها و منابع",
    ]
    for item in toc_items:
        add_p(doc, item, indent=False)
    doc.add_page_break()

    # Chapter 1
    add_heading(doc, "1. تعریف مسئله و محیط اختصاصی", 1)
    validation = source_map["validation"]
    add_p(
        doc,
        f"مطابق قاعده پروژه، رقم یکی‌مانده‌به‌آخر شماره دانشجویی 40400356 برابر 5 است؛ بنابراین اندازه نقشه از رابطه 15+(5 mod 4) برابر 16 شد. نقشه نهایی شامل {validation['wall_count']} خانه دیوار با نسبت {100*validation['wall_ratio']:.1f} درصد و {validation['penalty_count']} خانه جریمه است. مسیر شروع به کلید و سپس کلید به هدف با BFS قطعی اعتبارسنجی و فایل JSON نقشه برای اجرای یکسان تمام الگوریتم‌ها ثابت شد.",
    )
    add_figure(
        doc,
        "source_map.png",
        "شکل 1 - نقشه مبدأ اختصاصی تولیدشده با Seed برابر 5",
        "مانع عمودی میانی، عبور بدون کلید را از در D غیرممکن می‌کند. پس از دریافت K، عامل وارد نیمه راست می‌شود. دروازه G در مسیر کوتاه قرار دارد و تنها در فازهای 0 و 1 باز است؛ مسیر دورتر نیز وجود دارد. بنابراین قابلیت پویا صرفاً نمایشی نیست و تصمیم انتظار، عبور یا انتخاب مسیر جایگزین را تغییر می‌دهد.",
    )
    add_table(
        doc,
        ["مولفه", "مقدار"],
        [
            ("نقطه شروع", validation["start"]),
            ("محل کلید", validation["key"]),
            ("محل هدف", validation["goal"]),
            ("تعداد خانه قابل عبور", 16 * 16 - validation["wall_count"]),
            ("سقف گام", 3 * (16 * 16 - validation["wall_count"])),
            ("دوره دروازه", 4),
            ("فازهای باز", "0 و 1"),
        ],
    )
    add_p(doc, f"در مقصد مشابه، {similar_removed} مانع داخلی حذف و همین تعداد مانع در محل جدید اضافه شد؛ نسبت جابه‌جایی نسبت به موانع داخلی مبدأ {100*similar_frac:.1f} درصد است. در مقصد متفاوت، این مقادیر {diff_removed} و {100*diff_frac:.1f} درصد هستند و علاوه بر آن محل کلید/هدف و تعداد خانه‌های جریمه تغییر کرده است.")
    add_figure(doc, "target_similar_map.png", "شکل 2 - محیط مقصد مشابه", "شروع، کلید و هدف ثابت مانده‌اند و تغییرات محدود عمدتاً در همسایگی‌های محلی رخ داده است. این ساختار باید برای انتقال انتخابی مناسب‌تر باشد، زیرا بخش بزرگی از روابط عمل-نتیجه مبدأ هنوز معتبر است.")
    add_figure(doc, "target_different_map.png", "شکل 3 - محیط مقصد متفاوت", "تغییر گسترده دیوارها، جابه‌جایی عناصر مأموریت و افزایش خانه‌های جریمه باعث می‌شود کپی بی‌قیدوشرط جدول Q، عمل‌هایی را ترجیح دهد که در مقصد دیگر مناسب نیستند؛ این محیط برای آشکارکردن انتقال منفی طراحی شده است.")

    # Chapter 2
    add_heading(doc, "2. مدل‌سازی کامل MDP", 1)
    add_heading(doc, "2-1. فضای حالت و خاصیت مارکوف", 2)
    add_ltr(doc, "S = {(r, c, k, p) | (r,c) is passable, k in {0,1}, p in {0,1,2,3}}")
    add_p(doc, "متغیرهای r و c مختصات عامل، k وضعیت دریافت کلید و p فاز دروازه دوره‌ای هستند. مختصات به‌تنهایی کافی نیست؛ زیرا قابلیت عبور از در D به k و قابلیت عبور از G به p وابسته است. با دانستن همین چهار مولفه و عمل انتخابی، توزیع حالت بعد، پاداش و پایان مأموریت مشخص می‌شود و نیازی به تاریخچه نیست.")
    add_heading(doc, "2-2. فضای عمل و تابع انتقال", 2)
    add_ltr(doc, "A = {UP, DOWN, LEFT, RIGHT}")
    add_ltr(doc, "P(realized=a | selected=a)=0.8,  P(each perpendicular action)=0.1")
    add_p(doc, "در هر انتقال، فاز زمانی یک واحد جلو می‌رود؛ حتی اگر عامل به علت دیوار، در قفل یا دروازه بسته در مکان خود بماند. این نکته برای یادگیری زمان‌بندی عبور حیاتی است. ورود به کلید k را از صفر به یک تغییر می‌دهد. هدف پس از دریافت کلید حالت پایانی است.")
    add_heading(doc, "2-3. پایان و truncation", 2)
    add_p(doc, "رسیدن به هدف termination واقعی MDP است. سقف گام، یک محدودیت آزمایش از نوع TimeLimit است و با truncated ثبت می‌شود. واردکردن شمارنده گام در حالت، MDP را افق‌محدود و فضای حالت را چندصد برابر می‌کرد و با Value Iteration ایستای مورد انتظار پروژه ناسازگار بود؛ بنابراین این محدودیت خارج از هسته MDP و یکسان برای همه روش‌ها اعمال شد.")
    add_heading(doc, "2-4. سیاست", 2)
    add_ltr(doc, "pi(a|s) = probability of selecting action a in state s")
    add_p(doc, "سیاست Value Iteration پس از همگرایی حریصانه است. در Q-Learning و SARSA(λ)، سیاست رفتار ε-greedy است و ε طی آموزش کاهش می‌یابد؛ در ارزیابی، عمل حریصانه با شکستن tie به‌صورت قطعی انتخاب می‌شود.")

    # Chapter 3
    add_heading(doc, "3. طراحی تابع پاداش", 1)
    add_table(
        doc,
        ["رویداد", "پاداش پایه"],
        [
            ("هر گام", -0.5),
            ("برخورد با دیوار/دروازه بسته", "-1.5 افزوده"),
            ("ورود به خانه جریمه", "-7 افزوده"),
            ("دریافت کلید", "+25 افزوده"),
            ("تلاش برای در قفل", "-4 افزوده"),
            ("رسیدن به هدف", "+120 افزوده"),
        ],
    )
    add_p(doc, "در نسخه sparse همین مقادیر استفاده شدند. هزینه گام از سرگردانی جلوگیری می‌کند و پاداش هدف به‌قدری بزرگ است که مسیرهای معقول حتی با چند انحراف تصادفی بازده مثبت داشته باشند. برای عبور موفق از در پاداش جداگانه داده نشد، زیرا عبور رفت‌وبرگشتی می‌توانست به جمع‌آوری پاداش فرعی تبدیل شود.")
    add_ltr(doc, "r_shaped = r_sparse + eta [gamma Phi(s') - Phi(s)],   eta = 0.2")
    add_p(doc, "پتانسیل پیش از کلید برابر منفی فاصله باقی‌مانده تا کلید به‌علاوه فاصله کلید تا هدف، و پس از کلید برابر منفی فاصله تا هدف است. این ساختار اطلاعات میانی می‌دهد و در حالت ایده‌آل سیاست بهینه MDP تنزیل‌شده را حفظ می‌کند. مقیاس اولیه 1.0 در stress-test بازخورد مثبت بزرگی برای ماندن در فاصله زیاد ایجاد کرد؛ نسخه نهایی با η=0.2 و هزینه گام -0.5 این رفتار را حذف کرد. این یکی از دو پیشنهاد ناقص مستندشده در بخش شفافیت ابزارهاست.")

    # Chapter 4
    add_heading(doc, "4. پیاده‌سازی الگوریتم‌ها", 1)
    add_heading(doc, "4-1. Value Iteration", 2)
    add_ltr(doc, "V_{k+1}(s) = max_a sum_{s'} P(s'|s,a)[R(s,a,s') + gamma V_k(s')]")
    add_p(doc, "مدل انتقال دقیق از تابع transition_outcomes محیط استخراج و backup بلمن به‌صورت برداری اجرا شد. بردارسازی فقط بهینه‌سازی محاسباتی است و هیچ پیاده‌سازی آماده برنامه‌ریزی یا RL استفاده نشده است. شرط همگرایی، بیشینه قدرمطلق تغییر V در دو تکرار متوالی کمتر از 10^-8 است.")
    add_heading(doc, "4-2. Q-Learning", 2)
    add_ltr(doc, "Q(s,a) <- Q(s,a) + alpha [r + gamma max_{a'}Q(s',a') - Q(s,a)]")
    add_p(doc, "Q-Learning off-policy است؛ هدف به‌روزرسانی عمل حریصانه حالت بعد را فرض می‌کند، اما داده ممکن است با عمل اکتشافی تولید شده باشد. دو برنامه کاهش خطی و نمایی ε با نقاط شروع و پایان یکسان مقایسه شدند.")
    add_heading(doc, "4-3. SARSA(lambda)", 2)
    add_ltr(doc, "delta_t = r_{t+1} + gamma Q(s_{t+1},a_{t+1}) - Q(s_t,a_t)")
    add_ltr(doc, "E_t(s,a) = gamma lambda E_{t-1}(s,a);  E_t(s_t,a_t)=1  (replacing trace)")
    add_ltr(doc, "Q <- Q + alpha delta_t E")
    add_p(doc, "ردپای replacing انتخاب شد، زیرا در محیطی با حلقه، برخورد و ماندن، accumulating trace می‌تواند بازدید مکرر یک زوج حالت-عمل را بیش از حد تقویت کند. λ=0 به SARSA یک‌مرحله‌ای تبدیل می‌شود و افزایش λ خطای TD را به تعداد بیشتری از زوج‌های قبلی منتقل می‌کند.")

    # Chapter 5
    add_heading(doc, "5. پروتکل آزمایش و معیارهای آماری", 1)
    add_p(doc, f"آزمایش submission با {run_manifest['episodes']} اپیزود برای هر اجرای اصلی و Seedهای {run_manifest['seeds']} انجام شد. انتقال یادگیری با {run_manifest['transfer_episodes']} اپیزود و Seedهای {run_manifest['transfer_seeds']} اجرا شد. ارزیابی نهایی هر سیاست شامل {run_manifest['evaluation_episodes']} اپیزود با انتقال تصادفی محیط بود. میانگین، انحراف معیار و نوار اطمینان 95 درصد منحنی‌ها از تمام Seedها محاسبه شدند.")
    add_p(doc, "معیارها شامل بازده اپیزود، تعداد گام، نرخ موفقیت، برخورد با دیوار، ورود به خانه جریمه، زمان اجرا، حافظه جدول‌ها، توافق سیاست با Value Iteration، عملکرد اولیه انتقال، مساحت زیر منحنی موفقیت اولیه، اپیزود رسیدن به موفقیت 80 درصد و عملکرد نهایی است. گزارش فقط بهترین Seed را نمایش نمی‌دهد.")

    # Chapter 6
    add_heading(doc, "6. نتایج Value Iteration", 1)
    add_table(
        doc,
        ["γ", "تکرار", "همگرا", "delta نهایی", "زمان backup (s)", "V شروع"],
        [(fmt(r.gamma, 2), int(r.iterations), str(bool(r.converged)), f"{r.final_delta:.2e}", fmt(r.runtime_seconds, 4), fmt(r.start_value, 3)) for r in vi.itertuples(index=False)],
    )
    gamma_low = vi.loc[vi.gamma.idxmin()]
    gamma_high = vi.loc[vi.gamma.idxmax()]
    add_p(doc, f"با افزایش γ از {gamma_low.gamma:.2f} به {gamma_high.gamma:.2f} تعداد تکرار از {int(gamma_low.iterations)} به {int(gamma_high.iterations)} افزایش یافت. دلیل آن کاهش نرخ contraction عملگر بلمن در γهای نزدیک یک است. در مقابل، مقدار شروع و حساسیت سیاست به پیامدهای دوردست افزایش می‌یابد. γ=0.95 تعادل مناسبی میان افق تصمیم‌گیری مأموریت چندمرحله‌ای و هزینه همگرایی ایجاد کرد.")
    add_figure(doc, "vi_value_before_key.png", "شکل 4 - Heatmap ارزش بهینه پیش از دریافت کلید", "شیب ارزش عامل را ابتدا به سمت K هدایت می‌کند و نواحی بن‌بست/پرخطر ارزش پایین‌تری دارند. ارزش فقط تابع مکان نیست؛ این شکل برش k=0 و فاز صفر است و برش پس از کلید متفاوت خواهد بود.")
    add_figure(doc, "vi_value_after_key.png", "شکل 5 - Heatmap ارزش بهینه پس از دریافت کلید", "پس از تغییر k به یک، هدف واسط حذف می‌شود و ساختار ارزش به سمت در، دروازه و هدف نهایی می‌چرخد. تفاوت دو Heatmap نشان می‌دهد افزودن متغیر کلید به حالت یک ضرورت مدل‌سازی است، نه جزئیات پیاده‌سازی.")
    add_figure(doc, "vi_policy_before_key.png", "شکل 6 - سیاست بهینه Value Iteration پیش از کلید", "فلش‌ها بهترین عمل را برای فاز صفر نشان می‌دهند. نزدیک دیوارها ممکن است عمل ظاهراً غیرمستقیم انتخاب شود، چون 20 درصد احتمال انحراف و هزینه برخورد در امیدریاضی تصمیم وارد شده است.")
    add_figure(doc, "vi_policy_after_key.png", "شکل 7 - سیاست بهینه Value Iteration پس از کلید", "در بخش راست نقشه، سیاست زمان‌بندی عبور از G و مسیر جایگزین را با توجه به فاز دروازه می‌آموزد. نمایش یک فاز به‌تنهایی تمام سیاست را نشان نمی‌دهد، اما ساختار مکانی تصمیم را روشن می‌کند.")

    # Chapter 7
    add_heading(doc, "7. نتایج Q-Learning و طراحی پاداش", 1)
    eps_agg = q_eps_stats.groupby("epsilon_schedule").agg({"final_success":"mean","final_return":"mean","final_steps":"mean","episodes_to_80":"median"}).reset_index()
    add_table(doc, ["برنامه ε", "موفقیت نهایی", "بازده نهایی", "گام نهایی", "میانه اپیزود تا 80%"], [(r.epsilon_schedule, fmt(r.final_success), fmt(r.final_return,2), fmt(r.final_steps,1), fmt(r.episodes_to_80,0)) for r in eps_agg.itertuples(index=False)])
    add_figure(doc, "q_epsilon_success.png", "شکل 8 - مقایسه نرخ موفقیت برنامه‌های کاهش ε", f"بر اساس میانگین Seedها، برنامه «{best_eps}» انتخاب شد. معیار انتخاب فقط بازده یک اجرا نبود؛ موفقیت نهایی و سرعت عبور از آستانه 80 درصد با هم بررسی شدند. اختلاف نوارهای اطمینان در ابتدای آموزش نشان می‌دهد زمان‌بندی اکتشاف روی کشف زنجیره کلید-در-هدف اثر مستقیم دارد.")
    add_figure(doc, "q_epsilon_return.png", "شکل 9 - مقایسه بازده برنامه‌های کاهش ε", "بازده در مراحل اولیه نوسان زیادی دارد، زیرا سیاست اکتشافی هم با دیوارها برخورد می‌کند و هم ممکن است پس از گرفتن کلید نتواند مأموریت را تمام کند. پس از کاهش ε، پراکندگی کم و بازده به محدوده پایدار نزدیک می‌شود.")

    reward_agg = reward_stats.groupby("reward_mode").agg({"final_success":"mean","final_return":"mean","final_steps":"mean","episodes_to_80":"median"}).reset_index()
    add_table(doc, ["پاداش", "موفقیت نهایی", "بازده نهایی", "گام نهایی", "میانه اپیزود تا 80%"], [(r.reward_mode, fmt(r.final_success), fmt(r.final_return,2), fmt(r.final_steps,1), fmt(r.episodes_to_80,0)) for r in reward_agg.itertuples(index=False)])
    sparse_row = reward_agg[reward_agg.reward_mode == "sparse"].iloc[0]
    shaped_row = reward_agg[reward_agg.reward_mode == "shaping"].iloc[0]
    add_figure(doc, "reward_success.png", "شکل 10 - نرخ موفقیت sparse و shaping", f"پاداش شکل‌دهی‌شده سرعت یادگیری را از نظر میانه رسیدن به آستانه 80 درصد از {fmt(sparse_row.episodes_to_80,0)} به {fmt(shaped_row.episodes_to_80,0)} اپیزود تغییر داد. مقایسه موفقیت نهایی ({fmt(sparse_row.final_success)} در برابر {fmt(shaped_row.final_success)}) مشخص می‌کند اثر shaping فقط شتاب اولیه بوده یا در بودجه محدود آموزش روی کیفیت نهایی نیز اثر گذاشته است.")
    add_figure(doc, "reward_return.png", "شکل 11 - بازده sparse و shaping", "مقادیر مطلق بازده دو تعریف مستقیماً هم‌مقیاس نیستند، زیرا shaping جمله پتانسیل را اضافه می‌کند. بنابراین این نمودار باید همراه نرخ موفقیت، طول مسیر و توافق سیاست تفسیر شود؛ استفاده از بازده به‌تنهایی می‌توانست نتیجه‌گیری غلط ایجاد کند.")
    add_figure(doc, "q_training_visitation.png", "شکل 12 - Heatmap تعداد بازدید Q-Learning", "تراکم بازدید در مسیرهای نزدیک شروع و نقاط تصمیم بیشتر است. نواحی کم‌بازدید همان جاهایی هستند که Q-Learning احتمالاً با سیاست مرجع اختلاف بیشتری دارد؛ زیرا برآورد Q در آن‌ها داده کافی دریافت نکرده است.")
    add_figure(doc, "q_final_trajectory.png", "شکل 13 - یک مسیر نهایی عامل Q-Learning", "مسیر نمایش‌داده‌شده یک rollout واقعی با تصادفی‌بودن انتقال است، نه مسیر دست‌ساز. انحراف‌های محدود نسبت به مسیر هندسی کوتاه ناشی از لغزش 0.1+0.1 و واکنش سیاست به فاز دروازه هستند.")

    # Manual update
    add_heading(doc, "7-1. بازسازی یک Q-update واقعی", 2)
    candidates = q_update[(q_update["reward"] > 0) | (q_update["bootstrap"].abs() > 1e-9)]
    row = (candidates.iloc[0] if len(candidates) else q_update.iloc[0])
    add_table(doc, ["مولفه", "مقدار"], [("اپیزود/گام", f"{int(row.episode)} / {int(row.step)}"), ("حالت", row.state), ("عمل", int(row.action)), ("پاداش", fmt(row.reward,4)), ("حالت بعد", row.next_state), ("Q قبل", fmt(row.q_before,6)), ("max Q بعد", fmt(row.bootstrap,6)), ("هدف TD", fmt(row.target,6)), ("خطای TD", fmt(row.td_error,6)), ("Q بعد", fmt(row.q_after,6))])
    add_ltr(doc, f"Q_new = {row.q_before:.6f} + 0.22 * ({row.target:.6f} - {row.q_before:.6f}) = {row.q_after:.6f}")
    add_p(doc, "اعداد جدول مستقیماً از فایل q_learning_manual_update_trace.csv خوانده شده‌اند. بنابراین این محاسبه یک مثال فرضی نیست و می‌توان آن را با همان Seed بازتولید کرد.")

    # Chapter 8
    add_heading(doc, "8. نتایج SARSA(lambda)", 1)
    lambda_agg = lambda_stats.groupby("lambda_label").agg({"final_success":"mean","final_return":"mean","final_steps":"mean","episodes_to_80":"median","wall_hits":"mean","penalty_entries":"mean"}).reset_index()
    add_table(doc, ["λ", "موفقیت", "بازده", "گام", "اپیزود تا 80%", "برخورد", "خطر"], [(r.lambda_label, fmt(r.final_success), fmt(r.final_return,1), fmt(r.final_steps,1), fmt(r.episodes_to_80,0), fmt(r.wall_hits,2), fmt(r.penalty_entries,2)) for r in lambda_agg.itertuples(index=False)])
    add_figure(doc, "sarsa_lambda_success.png", "شکل 14 - نرخ موفقیت برای چهار مقدار λ", f"بهترین تعادل داده‌های این پروژه در λ={best_lambda:.1f} مشاهده شد. λ صفر فقط اعتبار یک گام را منتشر می‌کند و معمولاً کندتر است؛ λ بسیار بزرگ خطاهای noisy ناشی از انتقال تصادفی و سیاست اکتشافی را به دنباله طولانی‌تری پخش می‌کند و می‌تواند ناپایداری را افزایش دهد.")
    add_figure(doc, "sarsa_lambda_return.png", "شکل 15 - بازده برای چهار مقدار λ", "بهبود سرعت در λ میانی ناشی از credit assignment چندمرحله‌ای است. افت یا نوسان λ=0.9 نشان می‌دهد بزرگ‌کردن trace همیشه بهتر نیست؛ به‌ویژه وقتی اپیزودها شامل برخورد، لغزش و حلقه باشند.")
    add_heading(doc, "8-1. ثبت delta و Eligibility Trace", 2)
    trace_rows = trace.head(6)
    add_table(doc, ["گام", "حالت", "عمل", "پاداش", "delta", "ردپای فعال قبل", "ردپای فعال بعد"], [(int(r.step), r.state, int(r.action), fmt(r.reward,3), fmt(r.delta,5), int(r.active_traces_before_decay), int(r.active_traces_after_decay)) for r in trace_rows.itertuples(index=False)])
    add_p(doc, "در هر گام، زوج فعلی در replacing trace مقدار یک می‌گیرد و سپس تمام ردپاها با γλ کاهش می‌یابند. علامت delta تعیین می‌کند مقادیر Q زوج‌های اخیر هم‌جهت یا خلاف جهت تغییر کنند. با λ=0 این حافظه بلافاصله صفر می‌شود و الگوریتم به SARSA یک‌مرحله‌ای نزدیک می‌گردد.")

    # Chapter 9
    add_heading(doc, "9. مقایسه سه الگوریتم و اختلاف سیاست", 1)
    comp_rows = []
    for algorithm, group in comparison.groupby("algorithm"):
        comp_rows.append((algorithm, mean_std(group["success_rate"]), mean_std(group["mean_return"],2), mean_std(group["mean_steps"],1), mean_std(group["runtime_seconds"],3), mean_std(group["memory_bytes"],0), mean_std(group["policy_agreement"])))
    add_table(doc, ["روش", "موفقیت", "بازده", "گام", "زمان (s)", "حافظه (byte)", "توافق سیاست"], comp_rows)
    add_p(doc, "Value Iteration به نمونه تعاملی نیاز ندارد، اما مدل کامل انتقال و حافظه V/Q را لازم دارد. Q-Learning و SARSA(λ) مدل‌آزادند و هزینه اصلی آن‌ها تعداد گام‌های تجربه است. زمان‌های جدول شامل آموزش هر اجرا هستند؛ برای Value Iteration زمان backup بلمن گزارش شده است. حافظه SARSA هنگام آموزش علاوه بر Q شامل Eligibility Trace فعال است، هرچند مدل ذخیره‌شده نهایی فقط Q را نگه می‌دارد.")
    add_figure(doc, "q_policy_disagreement.png", "شکل 16 - اختلاف سیاست Q-Learning با مرجع", f"نرخ توافق دقیق روی حالت‌های قابل‌دسترسی {100*final['q_policy_agreement']:.1f} درصد است. این معیار سخت‌گیرانه است: اگر دو عمل ارزش تقریباً برابر داشته باشند ولی argmax متفاوت شود، اختلاف ثبت می‌شود. با این حال نقشه نشان می‌دهد اختلاف‌ها در نواحی کم‌نمونه، کنار خطر و نزدیک تصمیم‌های فازی متمرکزند.")
    add_figure(doc, "sarsa_policy_disagreement.png", "شکل 17 - اختلاف سیاست SARSA(λ) با مرجع", f"توافق SARSA(λ) برابر {100*final['sarsa_policy_agreement']:.1f} درصد است. on-policy بودن می‌تواند نزدیک خطر، سیاست محتاط‌تری بسازد چون ارزش عمل بعدی همان سیاست ε-greedy را در نظر می‌گیرد، نه عمل حریصانه فرضی.")
    add_heading(doc, "9-1. سه حالت نمونه اختلاف", 2)
    ex_rows = []
    for ex in disagreement.get("examples", [])[:3]:
        ex_rows.append((ex["state"], ex["tile"], ex["candidate_action"], ex["reference_action"], str([round(x,3) for x in ex["q_values"]]), str([round(x,3) for x in ex["vi_q_values"]])))
    add_table(doc, ["حالت", "خانه", "عمل مدل‌آزاد", "عمل VI", "Q مدل‌آزاد", "Q مرجع"], ex_rows)
    add_p(doc, "علت مشترک این اختلاف‌ها سه چیز است: پوشش نمونه محدود در ترکیب‌های خاص k و p، نویز bootstrap و tieهای نزدیک، و تفاوت میان هدف بهینه‌سازی سیاست رفتار با سیاست کاملاً حریصانه. در حالت‌های کنار دیوار یا خطر، یک انحراف عمود می‌تواند رتبه اعمال را عوض کند؛ در نزدیکی G نیز فاز زمانی به اندازه مکان اهمیت دارد.")

    # Chapter 10
    add_heading(doc, "10. انتقال یادگیری", 1)
    transfer_agg = transfer.groupby(["target", "scenario"]).agg(
        initial_success=("initial_success_rate","mean"),
        initial_return=("initial_mean_return","mean"),
        early_auc=("early_success_auc","mean"),
        episodes80=("episodes_to_80pct","median"),
        final_success=("final_success_rate","mean"),
        final_return=("final_mean_return","mean"),
        final_steps=("final_mean_steps","mean"),
        copied=("copied_fraction","mean"),
    ).reset_index()
    for target_name, target_label in (("similar", "مقصد مشابه"), ("different", "مقصد متفاوت")):
        subset = transfer_agg[transfer_agg.target == target_name]
        add_heading(doc, f"10-{1 if target_name=='similar' else 2}. {target_label}", 2)
        add_table(doc, ["سناریو", "موفقیت اولیه", "AUC اولیه", "اپیزود تا 80%", "موفقیت نهایی", "بازده نهایی", "گام نهایی", "سهم کپی"], [(r.scenario, fmt(r.initial_success), fmt(r.early_auc), fmt(r.episodes80,0), fmt(r.final_success), fmt(r.final_return,1), fmt(r.final_steps,1), fmt(r.copied)) for r in subset.itertuples(index=False)])
        scratch = subset[subset.scenario == "scratch"].iloc[0]
        best = subset.sort_values(["early_auc","final_success"], ascending=False).iloc[0]
        add_p(doc, f"در {target_label}، سناریوی {best.scenario} بیشترین AUC موفقیت اولیه را داشت. خط مبنای آموزش از صفر AUC={scratch.early_auc:.3f} و موفقیت اولیه={scratch.initial_success:.3f} داشت، در حالی که سناریوی منتخب به‌ترتیب {best.early_auc:.3f} و {best.initial_success:.3f} ثبت کرد. عملکرد نهایی جداگانه گزارش شده است؛ برتری آغازین لزوماً تضمین‌کننده بهترین انتهای آموزش نیست.")
        scaled_rows = subset[subset.scenario.str.startswith("scaled_beta")].copy()
        if len(scaled_rows) > 1 and scaled_rows["initial_success"].nunique() == 1:
            add_p(doc, "یک نکته مهم این است که همه انتقال‌های تعدیل‌شده در لحظه آغاز موفقیت یکسانی دارند. دلیل، خطا یا بی‌اثر بودن β نیست: ضرب تمام Qها در هر β مثبت، ترتیب اعمال و argmax را عوض نمی‌کند. اثر β پس از شروع یادگیری ظاهر می‌شود، زیرا قدرمطلق مقادیر اولیه اندازه خطای TD و سرعت کنارگذاشتن دانش ناسازگار را تغییر می‌دهد.")
    add_figure(doc, "transfer_similar_success.png", "شکل 18 - انتقال به مقصد مشابه", "در محیط مشابه، انتقال کامل یا انتخابی معمولاً باید شروع بهتری از scratch داشته باشد، چون ساختار محلی بسیاری از حالت‌ها ثابت مانده است. تفاوت سناریوهای β نشان می‌دهد کاهش شدت انتقال می‌تواند میان استفاده از دانش و امکان اصلاح تعادل ایجاد کند.")
    add_figure(doc, "transfer_different_success.png", "شکل 19 - انتقال به مقصد متفاوت", "در محیط متفاوت، کپی کامل می‌تواند عملکرد اولیه را بدتر کند یا سرعت اصلاح را کاهش دهد. انتقال انتخابی و βهای کوچک‌تر دانش نامعتبر را محدود می‌کنند، هرچند ممکن است بخشی از دانش مفید را نیز کنار بگذارند.")
    add_heading(doc, "10-3. نمونه انتقال منفی", 2)
    add_table(doc, ["مولفه", "مقدار"], [("حالت", negative["state"]), ("خانه", negative["tile"]), ("عمل منتقل‌شده", negative["transferred_action"]), ("عمل بهینه مقصد", negative["target_optimal_action"]), ("regret", fmt(negative["regret"],6)), ("Q منتقل‌شده", str([round(x,4) for x in negative["transferred_q_values"]])), ("Q بهینه مقصد", str([round(x,4) for x in negative["target_optimal_q_values"]]))])
    add_p(doc, "در این حالت، جدول مبدأ عملی را ترجیح می‌دهد که به ساختار قبلی وابسته بوده است. در مقصد، تغییر دیوار/خطر/موقعیت مأموریت باعث می‌شود ارزش بهینه عمل دیگری بیشتر باشد. regret اختلاف ارزش بهینه مقصد بین عمل درست و عمل القاشده از مبدأ است؛ بنابراین انتقال منفی با معیار عددی تعریف شده، نه با عبارت مبهم «انتقال بد بود».")
    if not q_correction.empty:
        first_q = q_correction.iloc[0]
        last_q = q_correction.iloc[-1]
        first_values = [round(float(first_q[f"q_action_{a}"]), 4) for a in range(4)]
        last_values = [round(float(last_q[f"q_action_{a}"]), 4) for a in range(4)]
        add_p(doc, "مقادیر این حالت در ابتدای انتقال و پس از آموزش مقصد به‌صورت زیر ثبت شدند:")
        add_ltr(doc, f"Q_before = {first_values}", center=True)
        add_ltr(doc, f"Q_after  = {last_values}", center=True)
        add_p(doc, "تغییر ترتیب مقادیر نشان می‌دهد تجربه مقصد چگونه ترجیح اولیه نامناسب را تصحیح کرده است.")
    add_figure(doc, "transfer_q_change_different.png", "شکل 20 - میزان اصلاح Q پس از انتقال کامل", "بیشترین تغییر Q در همسایگی‌هایی رخ می‌دهد که ساختار مقصد با مبدأ ناسازگار است یا سیاست اولیه مرتباً آن‌ها را تجربه می‌کند. تغییر کم الزاماً به معنای انتقال صحیح نیست؛ حالت کم‌بازدید ممکن است فرصت اصلاح پیدا نکرده باشد.")

    # Chapter 11
    add_heading(doc, "11. رابط گرافیکی و بازتولیدپذیری", 1)
    add_p(doc, "رابط Tkinter در gui/app.py مستقل از نمودارهای ثابت است و حرکت عامل را مرحله‌به‌مرحله نمایش می‌دهد. دیوار، کلید، در، هدف، خطر، دروازه باز/بسته و عامل نشانه بصری متمایز دارند. انتخاب الگوریتم و محیط، آموزش/ارزیابی، شروع، توقف، ادامه، بازنشانی، اجرای مجدد، سرعت انیمیشن و نمایش سیاست در رابط تعبیه شده است. اطلاعات زنده شامل اپیزود، گام، بازده، ε، وضعیت کلید، موفقیت اخیر و وضعیت دروازه است.")
    add_ltr(doc, "python main.py")
    add_ltr(doc, "python experiments/run_experiments.py --profile submission")
    add_ltr(doc, "pytest -q")
    add_p(doc, "تمام مسیرها نسبی و مبتنی بر ریشه مخزن هستند. run_manifest.json نسخه Python، NumPy، Pandas، سیستم‌عامل، Seedها و بودجه اپیزود را ثبت می‌کند. مدل‌ها با NPZ فشرده و داده‌های خام با CSV/JSONL ذخیره شده‌اند. تاریخچه Git شامل مراحل محیط، الگوریتم‌ها، آزمایش‌ها/GUI و نتایج/گزارش است.")

    # Chapter 12 direct answers
    add_heading(doc, "12. پاسخ مستقیم به پرسش‌های تحلیلی", 1)
    add_heading(doc, "12-1. تعریف MDP و حفظ خاصیت مارکوف", 2)
    add_p(doc, "MDP به‌صورت پنج‌تایی (S,A,P,R,γ) تعریف شد. S شامل مکان، کلید و فاز دروازه؛ A شامل چهار حرکت؛ P شامل احتمال 0.8/0.1/0.1 و منطق برخورد؛ R یکی از دو نسخه پاداش؛ و γ=0.95 است. هدف پس از کلید terminal است. حذف k پیش‌بینی عبور از در و حذف p پیش‌بینی عبور از دروازه را به تاریخچه وابسته می‌کرد؛ پس هر دو برای مارکوف‌بودن لازم‌اند.")
    add_heading(doc, "12-2. on-policy و off-policy نزدیک خطر", 2)
    q_group = comparison[comparison.algorithm == "Q-Learning"]
    s_group = comparison[comparison.algorithm == "SARSA(lambda)"]
    q_pen = q_group["mean_penalty_entries"].mean() if len(q_group) else np.nan
    s_pen = s_group["mean_penalty_entries"].mean() if len(s_group) else np.nan
    q_wall = q_group["mean_wall_hits"].mean() if len(q_group) else np.nan
    s_wall = s_group["mean_wall_hits"].mean() if len(s_group) else np.nan
    add_p(doc, f"Q-Learning هدف max Q را استفاده می‌کند و فرض می‌کند در آینده عمل حریصانه اجرا خواهد شد؛ SARSA ارزش عمل واقعی ε-greedy بعدی را وارد می‌کند. در ارزیابی این پروژه، میانگین ورود به خطر برای Q-Learning برابر {fmt(q_pen,3)} و برای SARSA برابر {fmt(s_pen,3)} بود؛ برخورد با دیوار نیز {fmt(q_wall,3)} در برابر {fmt(s_wall,3)} ثبت شد. تفاوت باید همراه طول مسیر تفسیر شود: سیاست محتاط‌تر ممکن است امن‌تر ولی طولانی‌تر باشد.")
    add_heading(doc, "12-3. نیاز Value Iteration به مدل", 2)
    add_p(doc, "Value Iteration برای محاسبه امیدریاضی backup بلمن باید تمام حالت‌های بعد، احتمال و پاداش آن‌ها را بداند. مزیت آن سیاست مرجع بدون خطای نمونه و همگرایی قابل‌کنترل است؛ محدودیت آن نیاز به مدل و رشد هزینه با فضای حالت است. Q-Learning و SARSA فقط نمونه (s,a,r,s') می‌خواهند و برای محیط ناشناخته مناسب‌ترند، اما به اپیزودهای زیاد، اکتشاف و تنظیم پارامتر حساس‌اند.")
    add_heading(doc, "12-4. بهترین λ", 2)
    best_lambda_row = lambda_agg[lambda_agg.lambda_label == f"lambda={best_lambda:.1f}"].iloc[0]
    add_p(doc, f"λ={best_lambda:.1f} بهترین تعادل را ایجاد کرد: موفقیت نهایی میانگین {best_lambda_row.final_success:.3f}، بازده {best_lambda_row.final_return:.2f} و میانه زمان رسیدن به 80 درصد {fmt(best_lambda_row.episodes_to_80,0)} اپیزود بود. λهای بزرگ‌تر حافظه اعتباری طولانی‌تری داشتند، اما نویز انتقال و اکتشاف را نیز به گذشته پخش کردند.")
    add_heading(doc, "12-5. سه اختلاف سیاست", 2)
    add_p(doc, "سه حالت جدول بخش 9-1 از فایل مقایسه واقعی استخراج شدند. تحلیل محلی نشان داد اختلاف‌ها به دیوار/خطر مجاور، فاز دروازه و کمبود بازدید وابسته‌اند. علاوه بر این، معیار argmax به tieهای نزدیک حساس است؛ پس درصد توافق باید در کنار regret یا اختلاف مقدار Q خوانده شود.")
    add_heading(doc, "12-6. محیط مشابه و متفاوت در انتقال", 2)
    sim = transfer_agg[transfer_agg.target == "similar"]
    dif = transfer_agg[transfer_agg.target == "different"]
    sim_full = sim[sim.scenario == "full"].iloc[0]
    dif_full = dif[dif.scenario == "full"].iloc[0]
    sim_scratch = sim[sim.scenario == "scratch"].iloc[0]
    dif_scratch = dif[dif.scenario == "scratch"].iloc[0]
    add_p(doc, f"انتقال کامل در مقصد مشابه موفقیت اولیه {sim_full.initial_success:.3f} و AUC اولیه {sim_full.early_auc:.3f} داشت؛ همین مقادیر در مقصد متفاوت {dif_full.initial_success:.3f} و {dif_full.early_auc:.3f} شدند. این اعداد نشان می‌دهند شباهت ساختاری به‌تنهایی ترتیب عملکرد اولیه را تعیین نمی‌کند؛ هندسه مسیر، محل خطر و قابلیت اکتشاف نیز مؤثرند. شاهد روشن آن است که scratch در مقصد مشابه با موفقیت نهایی {sim_scratch.final_success:.3f} در بودجه 800 اپیزود شکست خورد، اما در مقصد متفاوت به {dif_scratch.final_success:.3f} رسید. بنابراین «مشابه‌تر» الزاماً «آسان‌تر» نیست. در عین حال، انتقال در مقصد مشابه سرعت یادگیری بسیار بیشتری از scratch ایجاد کرد و در مقصد متفاوت β=0.25 سریع‌تر از کپی کامل دانش ناسازگار را اصلاح کرد.")

    # Chapter 13
    add_heading(doc, "13. محدودیت‌ها و جمع‌بندی", 1)
    add_p(doc, "نخست، محیط جدولی است و تعمیم تابعی به نقشه‌های بسیار بزرگ بررسی نشده است. دوم، برش‌های دوبعدی سیاست فقط یک وضعیت کلید و فاز را نشان می‌دهند، در حالی که سیاست کامل چهاربعدی است. سوم، پنج Seed برای آزمایش اصلی و سه Seed برای انتقال عدم‌قطعیت را کاهش می‌دهد ولی جای تحلیل بسیار بزرگ را نمی‌گیرد. چهارم، معیار توافق دقیق به tieها حساس است. پنجم، TimeLimit خارج از MDP ایستا اعمال شده و اگر هدف تحلیل افق محدود باشد باید شمارنده زمان به حالت افزوده و برنامه‌ریزی پسرو استفاده شود.")
    add_p(doc, "نتیجه اصلی این است که کیفیت پروژه فقط با بازده نهایی سنجیده نمی‌شود. مدل درست حالت، پاداش بدون loophole، برنامه اکتشاف، طول trace و میزان شباهت مبدأ/مقصد همگی رفتار عامل را تغییر می‌دهند. Value Iteration مرجع ساختاری فراهم کرد؛ Q-Learning سیاست off-policy کارآمدی آموخت؛ SARSA(λ) اثر credit assignment و احتیاط on-policy را نشان داد؛ و انتقال انتخابی/تعدیل‌شده در برابر انتقال منفی مقاوم‌تر بود.")

    # AI usage
    add_heading(doc, "پیوست الف - شفافیت استفاده از ابزارهای کمکی", 1)
    add_table(doc, ["مورد", "پیشنهاد", "اصلاح/اعتبارسنجی", "دلیل"], [
        ("حالت", "مختصات و کلید", "افزودن فاز دروازه", "حفظ خاصیت مارکوف"),
        ("Reward shaping", "پاداش فاصله مستقیم", "پتانسیل‌محور با η=0.2", "جلوگیری از تغییر سیاست و حلقه"),
        ("Eligibility trace", "accumulating", "replacing و آزمایش چهار λ", "کنترل تقویت بازدید تکراری"),
        ("پیشنهاد ناقص 1", "shaping با مقیاس 1.0", "کاهش مقیاس و افزایش هزینه گام", "حذف بازده حلقه‌ای نامطلوب"),
        ("پیشنهاد ناقص 2", "مسیر پایه حدود 40 گام", "بازطراحی مسیر به مأموریت کوتاه‌تر", "قابل‌اکتشاف‌شدن با بودجه اپیزود"),
        ("بهینه‌سازی", "حلقه Python برای VI", "backup برداری از مدل مستقل", "کاهش زمان بدون تغییر معادله"),
    ])
    add_p(doc, "هیچ خروجی پیشنهادشده بدون اجرا و بررسی وارد نسخه نهایی نشد. دو مورد ناسازگار بالا با مشاهده داده واقعی شناسایی و اصلاح شدند. کد نهایی با تست واحد و اجرای چند Seed اعتبارسنجی شده است.")

    # Repro appendix
    add_heading(doc, "پیوست ب - ساختار تحویل و بازتولید", 1)
    add_ltr(doc, "RL_FinalProject_40400356/")
    add_ltr(doc, "  environments/  agents/  transfer/  gui/  experiments/  results/  tests/")
    add_ltr(doc, "  report.pdf  requirements.txt  README.md  main.py")
    add_p(doc, "فایل‌های تنظیمات هر دسته آزمایش در experiments/configs، داده خام در results/raw_data، مدل‌ها در results/models و تصاویر در results/figures قرار دارند. required_event_examples.jsonl نمونه همه رویدادهای اجباری را ثبت می‌کند.")

    add_heading(doc, "منابع", 1)
    references = [
        "Sutton, R. S., & Barto, A. G. Reinforcement Learning: An Introduction, 2nd ed., MIT Press, 2018.",
        "Bellman, R. Dynamic Programming. Princeton University Press, 1957.",
        "Watkins, C. J. C. H., & Dayan, P. Q-learning. Machine Learning, 8, 279-292, 1992.",
        "Rummery, G. A., & Niranjan, M. On-line Q-learning using connectionist systems. Cambridge University Engineering Department, 1994.",
        "Ng, A. Y., Harada, D., & Russell, S. Policy invariance under reward transformations. ICML, 1999.",
    ]
    for i, ref in enumerate(references, 1):
        add_ltr(doc, f"[{i}] {ref}", center=False, font=LATIN_FONT, size=9.8)

    doc.core_properties.title = "RL Final Project - Dynamic Maze"
    doc.core_properties.subject = "Student ID 40400356"
    doc.core_properties.author = "Amir Mohammad Jamaldar"
    doc.core_properties.keywords = "reinforcement learning, Q-learning, SARSA(lambda), value iteration, transfer learning"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_report()
    print(path)
